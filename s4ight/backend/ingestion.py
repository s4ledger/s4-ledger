"""
S4ight document ingestion.

Per-session, in-memory ephemeral document store. Files are extracted to
text, chunked, embedded with the same OpenAI model used by the semantic
retriever, and made searchable for the duration of the session.

Why in-memory only:
  - Vercel serverless filesystems are read-only outside /tmp, and /tmp
    is per-instance and short-lived.
  - Ephemeral is also the right ATO default: nothing persists until you
    explicitly add a Supabase / S3 drain later.

Supported formats:
  - .pdf   (PyPDF2 if available)
  - .docx  (python-docx if available)
  - .xlsx  (openpyxl if available)
  - .txt / .md / .csv  (plain text)
"""

from __future__ import annotations

import io
import json
import logging
import math
import os
import re
import threading
import time
from typing import Any, Dict, List, Optional, Tuple

from config import OPENAI_API_KEY, OPENAI_BASE_URL

log = logging.getLogger("s4ight.ingestion")

EMBED_MODEL = os.getenv("S4IGHT_EMBED_MODEL", "text-embedding-3-small")
INGEST_MAX_BYTES = int(os.getenv("S4IGHT_INGEST_MAX_BYTES", str(8 * 1024 * 1024)))  # 8 MB
INGEST_MAX_CHARS = int(os.getenv("S4IGHT_INGEST_MAX_CHARS", str(400_000)))         # ~80–100 pages
INGEST_MAX_DOCS_PER_SESSION = int(os.getenv("S4IGHT_INGEST_MAX_DOCS", "20"))

# Allowed classification labels (informational; we don't enforce policy here).
ALLOWED_CLASSIFICATIONS = [
    "UNCLASSIFIED",
    "UNCLASSIFIED//FOUO",
    "CUI",
    "CUI//SP-PRVCY",
    "CUI//SP-PROCUREMENT",
    "PROPRIETARY",
]
DEFAULT_CLASSIFICATION = os.getenv("S4IGHT_DEFAULT_CLASSIFICATION", "UNCLASSIFIED")


def normalize_classification(label: Optional[str]) -> str:
    if not label:
        return DEFAULT_CLASSIFICATION
    lbl = str(label).strip().upper()
    return lbl if lbl else DEFAULT_CLASSIFICATION


def _cosine(a: List[float], b: List[float]) -> float:
    if not a or not b or len(a) != len(b):
        return 0.0
    dot = 0.0
    na = 0.0
    nb = 0.0
    for x, y in zip(a, b):
        dot += x * y
        na += x * x
        nb += y * y
    if na == 0.0 or nb == 0.0:
        return 0.0
    return dot / (math.sqrt(na) * math.sqrt(nb))


def _chunk_text(text: str, chunk_size: int = 900, overlap: int = 150) -> List[str]:
    if not text:
        return []
    chunks: List[str] = []
    start = 0
    n = len(text)
    while start < n:
        end = min(start + chunk_size, n)
        chunks.append(text[start:end])
        if end >= n:
            break
        start = end - overlap
    return chunks


_WS_RE = re.compile(r"[ \t]+")
_NL_RE = re.compile(r"\n{3,}")


def _normalize(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = _WS_RE.sub(" ", text)
    text = _NL_RE.sub("\n\n", text)
    return text.strip()


# ---------------- Extractors ----------------

def _extract_text(filename: str, content: bytes) -> str:
    ext = os.path.splitext(filename or "")[1].lower()
    if ext in (".txt", ".md", ".csv", ".log"):
        try:
            return content.decode("utf-8", errors="replace")
        except Exception:
            return content.decode("latin-1", errors="replace")

    if ext == ".pdf":
        try:
            from PyPDF2 import PdfReader  # type: ignore
        except Exception as e:
            raise RuntimeError(
                "PDF support requires PyPDF2. Install with `pip install PyPDF2`."
            ) from e
        reader = PdfReader(io.BytesIO(content))
        parts: List[str] = []
        for page in reader.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                continue
        return "\n\n".join(parts)

    if ext == ".docx":
        try:
            import docx  # python-docx  # type: ignore
        except Exception as e:
            raise RuntimeError(
                "DOCX support requires python-docx. Install with `pip install python-docx`."
            ) from e
        d = docx.Document(io.BytesIO(content))
        parts = [p.text for p in d.paragraphs if p.text]
        for table in d.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(parts)

    if ext == ".xlsx":
        try:
            from openpyxl import load_workbook  # type: ignore
        except Exception as e:
            raise RuntimeError(
                "XLSX support requires openpyxl. Install with `pip install openpyxl`."
            ) from e
        wb = load_workbook(io.BytesIO(content), data_only=True, read_only=True)
        parts: List[str] = []
        for sheet in wb.worksheets:
            parts.append(f"## Sheet: {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                cells = ["" if v is None else str(v) for v in row]
                if any(c.strip() for c in cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    raise RuntimeError(f"Unsupported file type: {ext or '(no extension)'}")


# ---------------- Store ----------------

class _DocChunk:
    __slots__ = ("source", "idx", "text", "embedding", "classification")

    def __init__(self, source: str, idx: int, text: str, embedding: Optional[List[float]] = None, classification: str = DEFAULT_CLASSIFICATION):
        self.source = source
        self.idx = idx
        self.text = text
        self.embedding = embedding or []
        self.classification = classification


class DocumentStore:
    """Session-scoped document store."""

    def __init__(self):
        self._lock = threading.RLock()
        # session_id -> { "docs": [meta], "chunks": [_DocChunk], "rehydrated": bool }
        self._sessions: Dict[str, Dict[str, Any]] = {}
        self._client = None

    def _rehydrate_if_needed(self, session_id: str) -> None:
        """If we don't have this session in memory yet, try to load it from Supabase."""
        if not session_id:
            return
        with self._lock:
            existing = self._sessions.get(session_id)
            if existing is not None:
                return  # already in memory
        # Fetch outside the lock to avoid blocking other sessions.
        try:
            from doc_persistence import enabled as _persist_on, fetch_session  # noqa: WPS433
            if not _persist_on():
                return
            data = fetch_session(session_id)
            if not data or not data.get("docs"):
                return
            docs_by_id: Dict[str, Dict[str, Any]] = {d["id"]: d for d in data["docs"]}
            doc_chunks_by_id: Dict[str, List[_DocChunk]] = {}
            for row in data.get("chunks", []):
                d = docs_by_id.get(row["doc_id"])
                if not d:
                    continue
                emb = row.get("embedding") or []
                if isinstance(emb, str):
                    try:
                        emb = json.loads(emb)
                    except Exception:
                        emb = []
                chunk = _DocChunk(
                    source=d["name"], idx=int(row["idx"]),
                    text=row["text"] or "",
                    embedding=emb,
                    classification=row.get("classification") or DEFAULT_CLASSIFICATION,
                )
                doc_chunks_by_id.setdefault(row["doc_id"], []).append(chunk)
            session_state = {"docs": [], "chunks": []}
            for d in data["docs"]:
                meta = {
                    "id": d["id"],
                    "name": d["name"],
                    "chars": d["chars"],
                    "chunks": d["chunk_count"],
                    "classification": d.get("classification") or DEFAULT_CLASSIFICATION,
                    "uploaded_at": d.get("uploaded_at"),
                }
                session_state["docs"].append(meta)
                session_state["chunks"].extend(doc_chunks_by_id.get(d["id"], []))
            with self._lock:
                # Only install if still empty (another thread may have raced).
                if session_id not in self._sessions:
                    self._sessions[session_id] = session_state
                    log.info(
                        "Rehydrated session %s from Supabase: %d docs / %d chunks",
                        session_id, len(session_state["docs"]), len(session_state["chunks"]),
                    )
        except Exception as e:  # pragma: no cover
            log.warning("Rehydrate failed for %s: %s", session_id, e)

    # --- Public API ---

    def info(self, session_id: str) -> Dict[str, Any]:
        self._rehydrate_if_needed(session_id)
        with self._lock:
            s = self._sessions.get(session_id) or {"docs": [], "chunks": []}
            return {
                "session_id": session_id,
                "doc_count": len(s["docs"]),
                "chunk_count": len(s["chunks"]),
                "allowed_classifications": ALLOWED_CLASSIFICATIONS,
                "default_classification": DEFAULT_CLASSIFICATION,
                "docs": [
                    {
                        "id": d["id"],
                        "name": d["name"],
                        "chars": d["chars"],
                        "chunks": d["chunks"],
                        "classification": d.get("classification", DEFAULT_CLASSIFICATION),
                        "uploaded_at": d["uploaded_at"],
                    }
                    for d in s["docs"]
                ],
            }

    def ingest(self, session_id: str, filename: str, content: bytes, classification: Optional[str] = None) -> Dict[str, Any]:
        if not OPENAI_API_KEY:
            raise RuntimeError("Ingestion requires OPENAI_API_KEY to compute embeddings.")
        if not content:
            raise RuntimeError("Empty file.")
        if len(content) > INGEST_MAX_BYTES:
            raise RuntimeError(f"File too large (>{INGEST_MAX_BYTES} bytes).")

        text = _normalize(_extract_text(filename, content))
        if not text:
            raise RuntimeError("Could not extract any text from this file.")
        if len(text) > INGEST_MAX_CHARS:
            text = text[:INGEST_MAX_CHARS]

        chunks = _chunk_text(text)
        if not chunks:
            raise RuntimeError("No usable text after chunking.")

        # Embed chunks
        vectors = self._embed(chunks)

        with self._lock:
            sess = self._sessions.setdefault(session_id, {"docs": [], "chunks": []})
            if len(sess["docs"]) >= INGEST_MAX_DOCS_PER_SESSION:
                raise RuntimeError(
                    f"Session document limit reached ({INGEST_MAX_DOCS_PER_SESSION})."
                )
            doc_id = f"doc_{int(time.time() * 1000)}_{len(sess['docs']) + 1}"
            cls = normalize_classification(classification)
            new_chunks: List[Dict[str, Any]] = []
            for i, (c, v) in enumerate(zip(chunks, vectors), start=1):
                sess["chunks"].append(_DocChunk(
                    source=filename, idx=i, text=c, embedding=v, classification=cls
                ))
                new_chunks.append({
                    "idx": i, "text": c, "embedding": v, "classification": cls,
                })
            meta = {
                "id": doc_id,
                "name": filename,
                "chars": len(text),
                "chunks": len(chunks),
                "classification": cls,
                "uploaded_at": time.time(),
            }
            sess["docs"].append(meta)

        # Best-effort Supabase persistence (no-op if env vars absent).
        try:
            from doc_persistence import persist_document  # noqa: WPS433
            persist_document(session_id, meta, new_chunks)
        except Exception:  # pragma: no cover
            pass

        return {"status": "ok", "document": meta, "session": self.info(session_id)}

    def remove(self, session_id: str, doc_id: str) -> Dict[str, Any]:
        with self._lock:
            sess = self._sessions.get(session_id)
            if not sess:
                return {"status": "ok", "removed": 0}
            doc = next((d for d in sess["docs"] if d["id"] == doc_id), None)
            if not doc:
                return {"status": "ok", "removed": 0}
            name = doc["name"]
            sess["docs"] = [d for d in sess["docs"] if d["id"] != doc_id]
            before = len(sess["chunks"])
            sess["chunks"] = [c for c in sess["chunks"] if c.source != name]
            try:
                from doc_persistence import delete_document  # noqa: WPS433
                delete_document(doc_id)
            except Exception:  # pragma: no cover
                pass
            return {
                "status": "ok",
                "removed": before - len(sess["chunks"]),
                "session": self.info(session_id),
            }

    def clear(self, session_id: str) -> Dict[str, Any]:
        with self._lock:
            self._sessions.pop(session_id, None)
        try:
            from doc_persistence import clear_session  # noqa: WPS433
            clear_session(session_id)
        except Exception:  # pragma: no cover
            pass
        return {"status": "ok", "cleared": True, "session_id": session_id}

    def has_documents(self, session_id: str) -> bool:
        self._rehydrate_if_needed(session_id)
        with self._lock:
            sess = self._sessions.get(session_id)
            return bool(sess and sess["chunks"])

    def search(self, session_id: str, query: str, top_k: int = 4, allowed_classifications: Optional[List[str]] = None) -> List[Tuple[float, str, int, str, str]]:
        if not query or not query.strip():
            return []
        self._rehydrate_if_needed(session_id)
        with self._lock:
            sess = self._sessions.get(session_id)
            if not sess or not sess["chunks"]:
                return []
            chunks_snapshot = list(sess["chunks"])
        try:
            qvec = self._embed([query])[0]
        except Exception as e:  # pragma: no cover
            log.warning("Doc search embed failed: %s", e)
            return []
        allow = (
            {c.strip().upper() for c in allowed_classifications}
            if allowed_classifications else None
        )
        scored: List[Tuple[float, str, int, str, str]] = []
        for c in chunks_snapshot:
            if allow is not None and c.classification.upper() not in allow:
                continue
            sim = _cosine(qvec, c.embedding)
            scored.append((sim, c.source, c.idx, c.text, c.classification))
        scored.sort(reverse=True, key=lambda x: x[0])
        return scored[:top_k]

    # --- Internals ---

    def _client_obj(self):
        if self._client is not None:
            return self._client
        from openai import OpenAI  # type: ignore
        kwargs = {"api_key": OPENAI_API_KEY, "timeout": 30}
        if OPENAI_BASE_URL:
            kwargs["base_url"] = OPENAI_BASE_URL
        self._client = OpenAI(**kwargs)
        return self._client

    def _embed(self, texts: List[str]) -> List[List[float]]:
        client = self._client_obj()
        # Single-batch is fine for our chunk count; batch later if we go bigger.
        resp = client.embeddings.create(model=EMBED_MODEL, input=texts)
        return [d.embedding for d in resp.data]


# Global store
store = DocumentStore()
